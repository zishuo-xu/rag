import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


def parse_file(file_path: str, file_type: str) -> tuple[str, dict]:
    path = Path(file_path)
    parser = {
        "txt": _parse_plain_text,
        "md": _parse_plain_text,
        "pdf": _parse_pdf,
        "docx": _parse_docx,
    }.get(file_type)

    if parser is None:
        raise ValueError(f"unsupported file type: {file_type}")

    return parser(path)


def _parse_plain_text(path: Path) -> tuple[str, dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return text, {}


def _parse_pdf(path: Path) -> tuple[str, dict]:
    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    extracted_pages: list[str] = []
    text_page_count = 0

    for index, page in enumerate(reader.pages, start=1):
        raw_text = page.extract_text() or ""
        normalized = _normalize_pdf_page_text(raw_text)
        if _is_meaningful_pdf_page(normalized):
            text_page_count += 1
            extracted_pages.append(f"## Page {index}\n\n{normalized}")

    if not extracted_pages:
        raise ValueError("pdf contains no extractable text layer")

    return "\n\n".join(extracted_pages), {
        "page_count": page_count,
        "text_page_count": text_page_count,
        "empty_page_count": page_count - text_page_count,
    }


def _parse_docx(path: Path) -> tuple[str, dict]:
    doc = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    return "\n".join(paragraphs), {"paragraph_count": len(doc.paragraphs)}


def _normalize_pdf_page_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _is_meaningful_pdf_page(text: str) -> bool:
    compact = text.strip()
    if len(compact) < 20:
        return False
    visible_chars = sum(1 for char in compact if char.isalnum() or "\u4e00" <= char <= "\u9fff")
    return visible_chars >= 20
